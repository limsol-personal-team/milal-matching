import datetime
import os
import subprocess
import tempfile
from pathlib import Path
from users.utils import calculate_volunteer_hours_with_bonus
from records.models import VolunteerHours

try:
    import docx
    from docx.shared import Pt
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


def check_libreoffice():
    """Check if LibreOffice is installed and available"""
    try:
        # Try common LibreOffice command names
        commands = ['libreoffice', 'soffice']
        for cmd in commands:
            try:
                subprocess.run([cmd, '--version'], 
                             capture_output=True, 
                             check=True, 
                             timeout=10)
                return True
            except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
                continue
        return False
    except Exception:
        return False


def convert_datetime_to_month_year(start_date, end_date):
    """Convert datetime strings to month and year strings"""
    start_month = datetime.datetime.strptime(start_date, '%Y-%m-%d').strftime('%Y-%m')
    end_month = datetime.datetime.strptime(end_date, '%Y-%m-%d').strftime('%Y-%m')
    return start_month, end_month


def generate_docx_report(volunteer, start_date, end_date, total_hours, output_file, template_path=None):
    """Generate a Word document report by filling in the template"""
    if not DOCX_SUPPORT:
        raise Exception('DOCX support is not available. Please install python-docx library.')
    
    # Find the template docx path if not provided
    if template_path is None:
        # Look in the management/commands directory
        current_dir = Path(__file__).parent
        template_path = current_dir / 'templates' / 'volunteer_report_template.docx'
    
    if not template_path.exists():
        raise Exception(f'Word template not found at {template_path}')
    
    # Get the current date
    current_date = datetime.datetime.now().strftime('%Y-%m-%d')
    
    # Prepare data for replacement
    volunteer_name = f"{volunteer.first_name} {volunteer.last_name}"
    
    # Format dates as YYYY-MM
    start_month, end_month = convert_datetime_to_month_year(start_date, end_date)
    volunteer_period = f"{start_month} to {end_month}"
    
    # Load the template
    doc = docx.Document(str(template_path))
    
    # Replace placeholders in the document
    replacements = {
        'CURRENT_DATE': current_date,
        'NAME': volunteer_name,
        'PERIOD': volunteer_period,
        'TOTAL_HOURS': str(total_hours)
    }
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                # Replace text while preserving formatting
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            # Replace text while preserving formatting
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, value)
    
    # Save the document
    doc.save(output_file)


def convert_docx_to_pdf_libreoffice(docx_file, pdf_file):
    """Convert a DOCX file to PDF using LibreOffice"""
    if not check_libreoffice():
        raise Exception('LibreOffice is not installed or not available')
    
    # Get absolute paths
    docx_path = os.path.abspath(docx_file)
    output_dir = os.path.dirname(os.path.abspath(pdf_file))
    
    # LibreOffice command for conversion
    commands = ['libreoffice', 'soffice']
    
    for cmd in commands:
        try:
            # Run LibreOffice conversion
            result = subprocess.run([
                cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                docx_path
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                # LibreOffice creates PDF with same name as DOCX but .pdf extension
                generated_pdf = os.path.join(output_dir, 
                                           os.path.splitext(os.path.basename(docx_file))[0] + '.pdf')
                
                # Rename to desired filename if different
                if generated_pdf != os.path.abspath(pdf_file):
                    if os.path.exists(generated_pdf):
                        os.rename(generated_pdf, os.path.abspath(pdf_file))
                
                return
            else:
                raise Exception(f'LibreOffice error: {result.stderr}')
                
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            continue
    
    raise Exception('Failed to convert DOCX to PDF. LibreOffice conversion failed.')


def generate_text_report(volunteer, hours_entries, total_hours, bonus_hours, start_date, end_date, output_file):
    """Generate a simple text report for the volunteer's hours"""
    with open(output_file, 'w') as f:
        f.write(f"VOLUNTEER HOURS REPORT\n")
        f.write(f"=====================\n\n")
        f.write(f"Volunteer: {volunteer.first_name} {volunteer.last_name}\n")
        f.write(f"Report Date: {datetime.datetime.now().strftime('%Y-%m-%d')}\n")
        f.write(f"Period: {start_date} to {end_date}\n")
        f.write(f"Total Hours: {total_hours}\n\n")
        f.write(f"Detailed Entries:\n")
        f.write(f"----------------\n\n")
        
        for entry in hours_entries:
            service_type_display = dict(VolunteerHours.SERVICE_TYPES).get(entry.service_type, 'Unknown')
            service_date_str = entry.service_date.strftime('%Y-%m-%d')
            
            f.write(f"Service Date: {service_date_str}\n")
            f.write(f"Type: {service_type_display}\n")
            f.write(f"Hours: {entry.hours}\n")
            f.write(f"\n")
        
        if bonus_hours > 0:
            f.write(f"Adjusted Hours:\n")
            f.write(f"----------------\n\n")
            f.write(f"Type: Student Lead/Team Member\n")
            f.write(f"Hours: {bonus_hours}\n")


def generate_hours_logs_file(volunteer, hours_entries, start_date, end_date):
    """Generate a text file with detailed volunteer hours logs"""
    
    # Calculate total hours including bonus percentage
    total_hours, bonus_hours = calculate_volunteer_hours_with_bonus(volunteer, hours_entries)
    
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(
        mode='w',
        suffix='.txt',
        delete=False,
        encoding='utf-8'
    )
    
    # Generate the text report using the same data as the PDF
    generate_text_report(volunteer, hours_entries, total_hours, bonus_hours, start_date, end_date, temp_file.name)
    
    return temp_file.name


def generate_volunteer_pdf_report(volunteer, hours_entries):
    """Generate a PDF report for the volunteer's hours"""
    # Calculate total hours including bonus percentage
    total_hours, bonus_hours = calculate_volunteer_hours_with_bonus(volunteer, hours_entries)
    
    # Get date range
    earliest_record = hours_entries.earliest('service_date')
    latest_record = hours_entries.latest('service_date')
    start_date = earliest_record.service_date.strftime('%Y-%m-%d')
    end_date = latest_record.service_date.strftime('%Y-%m-%d')
    
    # Create temporary file for PDF
    temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
    pdf_path = temp_pdf.name
    temp_pdf.close()
    
    if DOCX_SUPPORT:
        # Generate DOCX first, then convert to PDF
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        docx_path = temp_docx.name
        temp_docx.close()
        
        try:
            generate_docx_report(volunteer, start_date, end_date, total_hours, docx_path)
            convert_docx_to_pdf_libreoffice(docx_path, pdf_path)
            
            # Clean up DOCX file
            if os.path.exists(docx_path):
                os.remove(docx_path)
                
        except Exception as e:
            # If PDF generation fails, create a simple text report
            raise Exception(f'Failed to generate PDF report: {e}')
    else:
        raise Exception('DOCX support is not available')
    
    return pdf_path, start_date, end_date